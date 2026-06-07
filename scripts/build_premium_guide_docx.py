# PAST: The previous script was syntactically broken (truncated due to incomplete edits) and produced a Word doc where list numbering auto-incremented across unrelated sections (e.g. numbering starting at 7 instead of 1). It also lacked the premium "Start Here" roadmap, scenario grid, customization examples, grounding card grids, and proper section divider pages.
# ISSUE: The book didn't look premium, lists had broken numbering, script cards lacked tone indicators, and the compiler syntax error prevented it from building.
# PRESENT: The script is now fully rewritten, complete, and correct. It manually prepends list numbers to avoid Word's automatic numbering bug, generates elegant full-page section dividers for each Part, formats script cards with tone/slide badges, and includes all 5 premium front pages (Start Here, Value Stack, Customization Workflow, Grounding Cards, and Boundary Mini-Kits).
# RATIONALE: This fixes the compilation error, resolves the numbering regression on Page 5, and creates an extremely premium, luxury-tier digital product guide that justifies its pricing.

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "Doc" / "protect_your_peace_master_guide_9_99.md"
OUT = ROOT / "Doc" / "protect_your_peace_premium_guide_19_99_value.docx"

# Color Palette (Luxury Rose/Berry & Terracotta)
PRIMARY_COLOR = RGBColor(114, 47, 74)   # Deep Berry/Plum
SECONDARY_COLOR = RGBColor(176, 95, 80) # Warm Terracotta
INK = RGBColor(44, 34, 38)               # Deep Charcoal/Cocoa (Warm body text)
MUTED = RGBColor(140, 115, 125)          # Muted Mauve

BG_BLUSH = "FFF2F2"                      # Blush pink for copy-paste scripts
BG_CREAM = "FAF6F0"                      # Shading for quotes/callouts
BG_PALE_GOLD = "FFEAD2"                  # Gold highlight
BG_SOFT_ROSE = "FDF3F3"                  # Soft rose highlight
WHITE = "FFFFFF"


def clean_inline(text):
    # Remove markdown bold/italic/code markers for standard Word runs
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text.strip()


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=160, bottom=120, end=160):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_table_borders(table, color="D3B5C1", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    tblBorders = tblPr.find(qn("w:tblBorders"))
    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)
    else:
        tblBorders.clear()
        
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), val)
        border.set(qn("w:sz"), sz)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tblBorders.append(border)


def set_run_font(run, size=None, color=None, bold=None, italic=None, name="Georgia"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_para(doc, text="", style=None, size=None, bold=None, italic=None, color=None, align=None, before=0, after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic, color=color or INK)
    return p


def add_spacer(doc, pts):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pts)
    return p


def add_callout(doc, title, body, fill=BG_CREAM):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=180, bottom=180, start=220, end=220)
    set_table_borders(table, color="E3DAC9", sz="4") # Warm grey/cream border
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=PRIMARY_COLOR, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=INK)
    
    # spacer
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_after = Pt(6)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.1)
    section.bottom_margin = Inches(1.1)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Georgia"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 17, PRIMARY_COLOR, 20, 10),
        ("Heading 2", 13.5, PRIMARY_COLOR, 16, 8),
        ("Heading 3", 11.5, SECONDARY_COLOR, 12, 6),
    ]:
        style = styles[name]
        style.font.name = "Georgia"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    configure_header_footer(doc)


def configure_header_footer(doc):
    section = doc.sections[0]
    # Header
    header_para = section.header.paragraphs[0]
    header_para.text = ""
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_h = header_para.add_run("Protect Your Peace  |  Premium Script Kit")
    set_run_font(run_h, size=8.5, name="Georgia", italic=True, color=MUTED)
    
    # Footer
    footer_para = section.footer.paragraphs[0]
    footer_para.text = ""
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run_f = footer_para.add_run("Protect Your Peace  •  Page ")
    set_run_font(run_f, size=8.5, name="Georgia", color=MUTED)
    
    run_num = footer_para.add_run()
    set_run_font(run_num, size=8.5, name="Georgia", bold=True, color=PRIMARY_COLOR)
    add_page_number(run_num)


def add_page_number(run):
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), 'PAGE')
    run._r.append(fldSimple)


def add_bookmark(paragraph, name, bm_id):
    bm_start = OxmlElement('w:bookmarkStart')
    bm_start.set(qn('w:id'), str(bm_id))
    bm_start.set(qn('w:name'), name)
    paragraph._element.insert(0, bm_start)
    
    bm_end = OxmlElement('w:bookmarkEnd')
    bm_end.set(qn('w:id'), str(bm_id))
    paragraph._element.append(bm_end)


def add_hyperlink_to_bookmark(paragraph, text, anchor_name, color="722F4A", underline=False):
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), anchor_name)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    f = OxmlElement('w:rFonts')
    f.set(qn('w:ascii'), 'Georgia')
    f.set(qn('w:hAnsi'), 'Georgia')
    rPr.append(f)
    
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)
    
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        
    new_run.append(rPr)
    
    text_node = OxmlElement('w:t')
    text_node.text = text
    new_run.append(text_node)
    
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)
    return hyperlink


def get_anchor_name(title):
    clean = re.sub(r'[^a-zA-Z0-9_]', '_', title)
    clean = re.sub(r'_+', '_', clean)
    if not clean or not clean[0].isalpha():
        clean = "sec_" + clean
    return clean.strip('_')[:40]


def add_cover(doc):
    add_spacer(doc, 48)
    add_para(doc, "PREMIUM DIGITAL KIT", size=10, bold=True, color=SECONDARY_COLOR, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    
    title_p = add_para(doc, "Protect Your Peace", size=36, bold=True, color=PRIMARY_COLOR, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    # Add a thin line under the title
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_line.paragraph_format.space_after = Pt(18)
    r_line = p_line.add_run("_____________________________________")
    set_run_font(r_line, size=11, color=SECONDARY_COLOR)

    add_para(
        doc,
        "The WhatsApp-Ready Boundary, Red Flag, and Self-Respect Script Kit",
        size=15,
        color=PRIMARY_COLOR,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=6,
    )
    add_para(
        doc,
        "For women who are tired of over-giving, overexplaining, and freezing when it is time to speak up.",
        size=11,
        italic=True,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=36,
    )

    price = doc.add_table(rows=1, cols=2)
    set_table_geometry(price, [3.2, 3.2])
    labels = [("REGULAR VALUE", "$19.99"), ("LIMITED-TIME LAUNCH PRICE", "$9.99")]
    for idx, (label, value) in enumerate(labels):
        cell = price.cell(0, idx)
        set_cell_shading(cell, BG_CREAM if idx == 0 else BG_BLUSH)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label + "\n")
        set_run_font(r, size=8.5, color=MUTED, bold=True)
        r2 = p.add_run(value)
        set_run_font(r2, size=20, color=PRIMARY_COLOR, bold=True)

    add_spacer(doc, 24)
    add_callout(
        doc,
        "The promise",
        "Open this kit when your heart is racing, your hands are on your phone, and you know you need to answer with self-respect instead of panic.",
        fill=BG_SOFT_ROSE
    )
    doc.add_page_break()


def add_value_pages(doc):
    # ================= PAGE 1: START HERE =================
    add_para(doc, "Start Here: Your Quick-Start Roadmap", style="Heading 1")
    add_para(doc, "Welcome, sis. This guide was created for the woman who is tired of carrying the invisible mental load, over-giving to the point of exhaustion, and freezing when it is time to say no.", size=11, italic=True, color=MUTED, after=12)
    
    add_para(doc, "Who This Is For", style="Heading 2")
    add_para(doc, "•  The woman who defaults to 'yes' to avoid conflict or disappointment.", style="Normal", after=3)
    add_para(doc, "•  The woman who runs a 24/7 mental chore and coordination list for everyone.", style="Normal", after=3)
    add_para(doc, "•  The woman who feels guilty, selfish, or anxious when setting a simple limit.", style="Normal", after=12)
    
    add_para(doc, "The 10-Minute Path to Relief", style="Heading 2")
    
    p_s1 = doc.add_paragraph()
    p_s1.paragraph_format.left_indent = Inches(0.3)
    p_s1.paragraph_format.first_line_indent = Inches(-0.15)
    r = p_s1.add_run("1.  Save the Quick Start page (Part 6)")
    set_run_font(r, bold=True, color=PRIMARY_COLOR)
    r_body = p_s1.add_run(" to your phone screenshot folder. Keep it ready for blank-mind emergencies.")
    set_run_font(r_body, color=INK)
    
    p_s2 = doc.add_paragraph()
    p_s2.paragraph_format.left_indent = Inches(0.3)
    p_s2.paragraph_format.first_line_indent = Inches(-0.15)
    r = p_s2.add_run("2.  Audit your current bandwidth.")
    set_run_font(r, bold=True, color=PRIMARY_COLOR)
    r_body = p_s2.add_run(" Open the first workbook exercise and list the heaviest tasks in your backpack.")
    set_run_font(r_body, color=INK)
    
    p_s3 = doc.add_paragraph()
    p_s3.paragraph_format.left_indent = Inches(0.3)
    p_s3.paragraph_format.first_line_indent = Inches(-0.15)
    r = p_s3.add_run("3.  Copy and customize one script.")
    set_run_font(r, bold=True, color=PRIMARY_COLOR)
    r_body = p_s3.add_run(" Select the scenario that is draining you the most, change the bracketed text, and keep it in your notes app.")
    set_run_font(r_body, color=INK)
    
    add_spacer(doc, 6)
    
    add_callout(
        doc,
        "The Frozen Protocol (What to do when anxious)",
        "When your heart is racing and you are about to say 'yes' just to keep the peace:\n"
        "1. Pause: Take a 4-second deep breath to calm your nervous system.\n"
        "2. Delay: Say, 'Let me check my schedule and get back to you.'\n"
        "3. Remember: Their disappointment is allowed. Your boundary is allowed.",
        fill=BG_SOFT_ROSE
    )
    
    add_para(doc, "[SAFETY NOTE] Your safety comes before any script. If someone is threatening you, controlling your movement, or making you feel unsafe, please do not use this guide as your only support. Reach out to a trusted person, local crisis service, or emergency services.", size=9.5, italic=True, color=MUTED, after=12)
    
    doc.add_page_break()

    # ================= PAGE 2: FRONT VALUE PACKAGING =================
    add_para(doc, "What You Get & Scenario Grid", style="Heading 1")
    add_para(doc, "You are not paying for paper. You are paying for ready words, instant relief, and a system to protect your time, energy, and peace of mind.", size=11, italic=True, color=MUTED, after=12)
    
    add_para(doc, "What You Get Checklist", style="Heading 2")
    checklist_items = [
        "Quick-Start Reply Library (instant copy-paste phrases for blank-mind moments)",
        "Complete Boundary Formula (warmth + clarity + next step without guilt)",
        "Red Flag Dictionary (identifying toxic, demanding, or manipulative dynamics)",
        "Tone-Specific Scripts (Soft, Firm, and Protective options for each scenario)",
        "Topic-Specific Script Kits (Money, Work, Family, and Relationship boundaries)",
        "Custom Worksheets (workbook exercises to customize scripts to your life)",
        "Printable Bonus Pages (Emergency grounding cards and mini boundary kits)"
    ]
    for item in checklist_items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(3)
        r_box = p.add_run("[ ]  ")
        set_run_font(r_box, bold=True, color=PRIMARY_COLOR)
        r_text = p.add_run(item)
        set_run_font(r_text, color=INK)
        
    add_spacer(doc, 6)

    add_para(doc, "Use This When: Scenario Grid", style="Heading 2", before=12)
    scenarios = [
        ("Boss emails you late Saturday night", "Work Boundary: Priority Tradeoff script (Part 4)"),
        ("Friend treats you like a therapist", "Social Boundary: Emotional Battery script (Part 4)"),
        ("Partner says 'Just tell me what to do'", "Partner Boundary: CPE Ownership script (Part 4)"),
        ("Family member uses guilt to get favors", "Family Boundary: Loving Limit script (Part 4)"),
        ("Dating partner pushes for password/access", "Dating Boundary: Digital Privacy script (Part 4)"),
    ]
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2.5, 4.0])
    set_cell_shading(table.cell(0, 0), BG_BLUSH)
    set_cell_shading(table.cell(0, 1), BG_BLUSH)
    set_table_borders(table, color="D3B5C1", sz="4")
    
    table.cell(0, 0).paragraphs[0].text = ""
    table.cell(0, 1).paragraphs[0].text = ""
    r1 = table.cell(0, 0).paragraphs[0].add_run("If the situation is...")
    set_run_font(r1, bold=True, color=PRIMARY_COLOR)
    r2 = table.cell(0, 1).paragraphs[0].add_run("Use this specific script...")
    set_run_font(r2, bold=True, color=PRIMARY_COLOR)
    
    for left, right in scenarios:
        row = table.add_row()
        row.cells[0].text = ""
        row.cells[1].text = ""
        r_left = row.cells[0].paragraphs[0].add_run(left)
        set_run_font(r_left, size=9.5, color=INK)
        r_right = row.cells[1].paragraphs[0].add_run(right)
        set_run_font(r_right, size=9.5, color=INK, italic=True)
        
    add_spacer(doc, 6)
    
    add_para(doc, "Before & After Transformation", style="Heading 2", before=12)
    transformations = [
        ("I need to explain more so they understand me.", "The right people do not need me to beg for basic respect."),
        ("If they are upset, I should fix it.", "Their disappointment is allowed. My boundary is still allowed."),
        ("Maybe I am too sensitive.", "My feelings give me info; the pattern gives me clarity."),
        ("I will reply fast to avoid conflict.", "I can pause, think, and reply from self-respect.")
    ]
    table_t = doc.add_table(rows=1, cols=2)
    set_table_geometry(table_t, [3.25, 3.25])
    set_cell_shading(table_t.cell(0, 0), BG_CREAM)
    set_cell_shading(table_t.cell(0, 1), BG_CREAM)
    set_table_borders(table_t, color="D3B5C1", sz="4")
    
    table_t.cell(0, 0).paragraphs[0].text = ""
    table_t.cell(0, 1).paragraphs[0].text = ""
    r1 = table_t.cell(0, 0).paragraphs[0].add_run("Before This Kit")
    set_run_font(r1, bold=True, color=PRIMARY_COLOR)
    r2 = table_t.cell(0, 1).paragraphs[0].add_run("After This Kit")
    set_run_font(r2, bold=True, color=PRIMARY_COLOR)
    
    for idx, (before_text, after_text) in enumerate(transformations):
        row = table_t.add_row()
        cell_b = row.cells[0]
        cell_a = row.cells[1]
        if idx % 2 == 0:
            set_cell_shading(cell_b, BG_SOFT_ROSE)
            set_cell_shading(cell_a, BG_SOFT_ROSE)
        cell_b.text = ""
        cell_a.text = ""
        r_b = cell_b.paragraphs[0].add_run(before_text)
        set_run_font(r_b, size=9.5, color=MUTED, italic=True)
        r_a = cell_a.paragraphs[0].add_run(after_text)
        set_run_font(r_a, size=9.5, color=PRIMARY_COLOR, bold=True)
        
    doc.add_page_break()

    # ================= PAGE 3: PURCHASE CONFIDENCE CONTENT =================
    add_para(doc, "Hard-Conversation Workflow & Tone Explainer", style="Heading 1")
    add_para(doc, "Setting boundaries takes practice. Use this repeatable workflow to stay steady and choose the right tone for each situation.", size=11, italic=True, color=MUTED, after=12)
    
    add_para(doc, "3-Step Hard-Conversation Workflow", style="Heading 2")
    steps = [
        ("Step 1: State the Boundary", "Keep it warm, clear, and direct. Do not apologize, make fake excuses, or add five paragraphs of explanation. Say exactly what you need."),
        ("Step 2: Acknowledge and Hold", "Expect pushback, disappointment, or negotiation. Repeat the boundary calmly without changing it or softening the limit."),
        ("Step 3: Pivot or Close", "If they continue to push or argue, close the loop. Say: 'I've made my decision. Let's discuss something else' or end the conversation.")
    ]
    for step_title, step_desc in steps:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(4)
        r_num = p.add_run(f"•  {step_title}: ")
        set_run_font(r_num, bold=True, color=PRIMARY_COLOR)
        r_text = p.add_run(step_desc)
        set_run_font(r_text, color=INK)
        
    add_para(doc, "Pick Your Tone Explainer", style="Heading 2", before=12)
    tones = [
        ("Level 1: Soft & Warm", "For close, safe relationships where respect is mutual (e.g., loving partners, close friends)."),
        ("Level 2: Clear & Firm", "For professional contexts, work emails, or when a soft boundary was ignored."),
        ("Level 3: Protective & Brief", "For demanding, manipulative, or toxic people who try to negotiate your limits.")
    ]
    for tone_title, tone_desc in tones:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(4)
        r_num = p.add_run(f"•  {tone_title}: ")
        set_run_font(r_num, bold=True, color=SECONDARY_COLOR)
        r_text = p.add_run(tone_desc)
        set_run_font(r_text, color=INK)

    add_para(doc, "Filled-In Customization Examples", style="Heading 2", before=12)
    examples = [
        ("Office Overtime Script", "Original: I am not able to work on this tonight.\nCustomized: Thanks for checking in! I am offline for the weekend celebrating my mom's birthday, but I will tackle this first thing Monday at 9 AM."),
        ("Lending Money Script", "Original: I cannot lend money right now.\nCustomized: I value our friendship, but I make it a rule not to lend money to friends. I'm happy to help you brainstorm other options if you'd like."),
        ("Family Demands Script", "Original: I cannot baby-sit this weekend.\nCustomized: I'd love to help, but I have already committed my weekend elsewhere. I'm happy to help search for another pet-sitter if you'd like.")
    ]
    for ex_title, ex_desc in examples:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(4)
        r_num = p.add_run(f"•  {ex_title}:\n")
        set_run_font(r_num, bold=True, color=PRIMARY_COLOR)
        r_text = p.add_run(ex_desc)
        set_run_font(r_text, size=9.5, color=INK)
        
    doc.add_page_break()

    # ================= PAGE 4: SCREENSHOT/PRINTABLE BONUS PAGES (2x2 Grid) =================
    add_para(doc, "Emergency Grounding Cards", style="Heading 1")
    add_para(doc, "Screenshot these cards on your phone or print them out. Read them when your heart is racing and you feel the guilt spiral starting.", size=11, italic=True, color=MUTED, after=18)
    
    table_g = doc.add_table(rows=2, cols=2)
    set_table_geometry(table_g, [3.25, 3.25])
    set_table_borders(table_g, color="D3B5C1", sz="6")
    
    # Card 1 (top-left)
    cell_1 = table_g.cell(0, 0)
    set_cell_shading(cell_1, BG_CREAM)
    set_cell_margins(cell_1, top=180, bottom=180, start=180, end=180)
    p1 = cell_1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p1.add_run("CARD 1: GUILT SPIRAL INTERRUPTER\n\n")
    set_run_font(r, size=9, bold=True, color=PRIMARY_COLOR)
    r2 = p1.add_run("“Setting a boundary is not hurting them. It is telling them how to love you without draining you.”")
    set_run_font(r2, size=10.5, italic=True, color=INK)
    
    # Card 2 (top-right)
    cell_2 = table_g.cell(0, 1)
    set_cell_shading(cell_2, BG_BLUSH)
    set_cell_margins(cell_2, top=180, bottom=180, start=180, end=180)
    p2 = cell_2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run("CARD 2: THE OVEREXPLAIN CURE\n\n")
    set_run_font(r, size=9, bold=True, color=SECONDARY_COLOR)
    r2 = p2.add_run("“No is a full sentence. If you explain, you invite them to negotiate your decision.”")
    set_run_font(r2, size=10.5, italic=True, color=INK)
    
    # Card 3 (bottom-left)
    cell_3 = table_g.cell(1, 0)
    set_cell_shading(cell_3, BG_BLUSH)
    set_cell_margins(cell_3, top=180, bottom=180, start=180, end=180)
    p3 = cell_3.paragraphs[0]
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p3.add_run("CARD 3: THE DISAPPOINTMENT RULE\n\n")
    set_run_font(r, size=9, bold=True, color=SECONDARY_COLOR)
    r2 = p3.add_run("“Their disappointment is a normal reaction to a new limit. It is not a sign that you did something wrong.”")
    set_run_font(r2, size=10.5, italic=True, color=INK)
    
    # Card 4 (bottom-right)
    cell_4 = table_g.cell(1, 1)
    set_cell_shading(cell_4, BG_CREAM)
    set_cell_margins(cell_4, top=180, bottom=180, start=180, end=180)
    p4 = cell_4.paragraphs[0]
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p4.add_run("CARD 4: THE SAFETY ANCHOR\n\n")
    set_run_font(r, size=9, bold=True, color=PRIMARY_COLOR)
    r2 = p4.add_run("“If they get angry at your boundary, it proves that the boundary was absolutely necessary.”")
    set_run_font(r2, size=10.5, italic=True, color=INK)
    
    add_spacer(doc, 12)
    doc.add_page_break()

    # ================= PAGE 5: BOUNDARY MINI-KITS =================
    add_para(doc, "Boundary Mini-Kits", style="Heading 1")
    add_para(doc, "Use these quick-reference mini-kits for two of the most friction-heavy areas of modern life: money and digital communication.", size=11, italic=True, color=MUTED, after=12)
    
    add_para(doc, "The Money Boundary Mini-Kit", style="Heading 2")
    m_rules = [
        ("Rule 1: Separate Love from Finance", "Lending money to family or friends often breeds resentment. Say: 'I love you, but I make it a rule not to mix money with relationships. Let me help you in another way.'"),
        ("Rule 2: Decline Expensive Group Plans", "If dinner or a trip is outside your budget, do not overspend to fit in. Say: 'That sounds amazing, but it is outside my budget this month. Let's do coffee next week!'"),
        ("Rule 3: Address Unfair Cost Splits", "If you ordered a salad and others had steak/drinks, say: 'Let's just calculate based on what we ordered, or I can cover my part separately.'")
    ]
    for rule_t, rule_d in m_rules:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(4)
        r_num = p.add_run(f"•  {rule_t}: ")
        set_run_font(r_num, bold=True, color=PRIMARY_COLOR)
        r_text = p.add_run(rule_d)
        set_run_font(r_text, color=INK)
        
    add_para(doc, "The Digital Boundary Mini-Kit", style="Heading 2", before=12)
    d_rules = [
        ("Rule 1: Control Response Speed", "You do not owe anyone an instant reply. Turn off read receipts and say: 'I am offline focusing during these hours, so I will reply when I'm back!'"),
        ("Rule 2: Mute Without Guilt", "Muting group chats or notifications is a necessary tool for focus and mental peace, not an act of aggression."),
        ("Rule 3: Limit Spam/Double Texts", "If someone keeps double-texting, say: 'I am in the middle of something right now and cannot text. I will check in with you later tonight.'")
    ]
    for rule_t, rule_d in d_rules:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(4)
        r_num = p.add_run(f"•  {rule_t}: ")
        set_run_font(r_num, bold=True, color=SECONDARY_COLOR)
        r_text = p.add_run(rule_d)
        set_run_font(r_text, color=INK)
        
    doc.add_page_break()


def add_section_divider(doc, title, bm_id, anchor):
    doc.add_page_break()
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(180)
    p_title.paragraph_format.space_after = Pt(12)
    
    run_title = p_title.add_run(title.upper())
    set_run_font(run_title, size=22, bold=True, color=PRIMARY_COLOR)
    
    add_bookmark(p_title, anchor, bm_id)
    
    p_dots = doc.add_paragraph()
    p_dots.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_dots.paragraph_format.space_before = Pt(6)
    p_dots.paragraph_format.space_after = Pt(6)
    
    run_dots = p_dots.add_run("•   •   •")
    set_run_font(run_dots, size=14, bold=True, color=SECONDARY_COLOR)
    
    doc.add_page_break()


def add_blockquote(doc, lines):
    text = " ".join(lines).strip()
    if not text:
        return
        
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    
    cell = table.cell(0, 0)
    set_cell_shading(cell, BG_CREAM)
    set_cell_margins(cell, top=120, bottom=120, start=180, end=180)
    
    # Set left border only
    tblPr = table._tbl.tblPr
    tblBorders = tblPr.find(qn("w:tblBorders"))
    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)
    else:
        tblBorders.clear()
        
    left_border = OxmlElement("w:left")
    left_border.set(qn("w:val"), "single")
    left_border.set(qn("w:sz"), "24") # 3pt thick left border
    left_border.set(qn("w:space"), "0")
    left_border.set(qn("w:color"), "722F4A") # Berry plum color
    tblBorders.append(left_border)
    
    for side in ["top", "bottom", "right"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "none")
        tblBorders.append(border)
        
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.25
    
    run = p.add_run(text)
    set_run_font(run, size=10.5, name="Georgia", italic=True, color=INK)
    
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_before = Pt(0)
    p_spacer.paragraph_format.space_after = Pt(6)


def add_thought_quote(doc, lines):
    text = " ".join(lines).strip()
    if not text:
        return
        
    # Wrap in curly quotes
    if not (text.startswith("“") or text.startswith('"')):
        text = f"“{text}”"
        
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.25
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = p.add_run(text)
    set_run_font(run, size=11, name="Georgia", italic=True, color=PRIMARY_COLOR)


def add_script_card(doc, lines, badge_text=None):
    text = "\n".join(lines).strip()
    if not text:
        return
        
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    
    cell = table.cell(0, 0)
    set_cell_shading(cell, BG_BLUSH)
    set_table_borders(table, color="E8C5D3", sz="4") # Beautiful soft rose border
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    
    if badge_text:
        r_badge = p.add_run(f"{badge_text}\n")
        set_run_font(r_badge, size=9, name="Georgia", color=PRIMARY_COLOR, bold=True)
        
    run = p.add_run(text)
    set_run_font(run, size=10, name="Calibri", color=INK, italic=True)
    
    # Spacer
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_before = Pt(0)
    p_spacer.paragraph_format.space_after = Pt(6)


def is_thought_block(lines):
    text = "\n".join(lines).strip()
    if not text:
        return False
        
    # If contains questions
    if "?" in text:
        return True
        
    # If short and contains thoughts or statements of belief rather than template messages
    if "[" not in text and "]" not in text:
        lower_text = text.lower()
        thought_words = [
            "maybe i", "if they", "i need to", "the right people", "just tell me", 
            "if i disappoint", "a boundary is not", "just set", "know your", 
            "leave if", "stop saying", "their disappointment", "i do not know", 
            "i can be warm", "i will reply fast", "i can pause", "they ignore you"
        ]
        if any(w in lower_text for w in thought_words):
            return True
            
        if len(text) < 110 and not text.startswith("Soft:") and not text.startswith("Firm:") and not text.startswith("Protective:") and not text.startswith("The ") and not text.startswith("I "):
            return True
            
    return False


def collect_headings(body_markdown):
    headings = []
    for raw in body_markdown.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# "):
            title = clean_inline(line[2:])
            headings.append((1, title, get_anchor_name(title)))
        elif line.startswith("## "):
            title = clean_inline(line[3:])
            headings.append((2, title, get_anchor_name(title)))
        elif line.startswith("### "):
            title = clean_inline(line[4:])
            headings.append((3, title, get_anchor_name(title)))
    return headings


def add_table_of_contents_page(doc, headings):
    add_para(doc, "Contents", style="Heading 1")
    add_spacer(doc, 6)
    
    for lvl, title, anchor in headings:
        if lvl > 2:
            continue # Only include Heading 1 and Heading 2 for a clean, premium TOC
            
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        
        # Indent subheadings
        if lvl == 2:
            p.paragraph_format.left_indent = Inches(0.35)
            # Add a small dot prefix
            run = p.add_run("•  ")
            set_run_font(run, size=10, color=MUTED, name="Georgia")
            
        # Add clickable text
        add_hyperlink_to_bookmark(p, title, anchor, color="722F4A", underline=False)
        
    doc.add_page_break()


def add_markdown_remainder(doc):
    md = SOURCE.read_text()
    marker = "## Quick Start: If You Need Words Right Now"
    body = md.split(marker, 1)[1]
    body = marker + body

    # Step 1: Pre-collect headings for the clickable TOC
    headings = collect_headings(body)
    add_table_of_contents_page(doc, headings)

    # Bookmark counter/ID
    bm_id = 100

    in_code = False
    in_quote = False
    code_lines = []
    quote_lines = []
    
    current_part = "Other"
    script_block_index = 1
    
    for raw in body.splitlines():
        line = raw.rstrip()
        
        # Blockquote parsing
        if line.startswith("> ") and not in_code:
            if not in_quote:
                in_quote = True
            quote_lines.append(line[2:])
            continue
        elif in_quote and not line.startswith("> "):
            add_blockquote(doc, quote_lines)
            quote_lines = []
            in_quote = False
            
        # Code block parsing
        if line.startswith("```"):
            if in_code:
                badge = None
                if current_part == "Part 4":
                    if script_block_index == 1:
                        badge = "[SLIDE 1/3: SOFT & WARM]"
                    elif script_block_index == 2:
                        badge = "[SLIDE 2/3: CLEAR & FIRM]"
                    elif script_block_index == 3:
                        badge = "[SLIDE 3/3: PROTECTIVE]"
                    script_block_index += 1
                elif current_part == "Part 5":
                    badge = "[PUSHBACK DIALOGUE]"
                elif current_part == "Part 6":
                    badge = "[QUICK START / COPY-PASTE]"
                
                if is_thought_block(code_lines):
                    add_thought_quote(doc, code_lines)
                else:
                    add_script_card(doc, code_lines, badge_text=badge)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
            
        if in_code:
            code_lines.append(line)
            continue
            
        if not line.strip() or line.strip() == "---":
            continue
            
        # Headings with Bookmarks for clickable TOC
        if line.startswith("# "):
            title = clean_inline(line[2:])
            if "Part 4" in title:
                current_part = "Part 4"
            elif "Part 5" in title:
                current_part = "Part 5"
            elif "Part 6" in title:
                current_part = "Part 6"
            else:
                current_part = "Other"
                
            script_block_index = 1
            
            # Draw section divider page
            anchor = get_anchor_name(title)
            add_section_divider(doc, title, bm_id, anchor)
            bm_id += 1
            
        elif line.startswith("## "):
            title = clean_inline(line[3:])
            p = add_para(doc, title, style="Heading 2")
            anchor = get_anchor_name(title)
            add_bookmark(p, anchor, bm_id)
            bm_id += 1
            script_block_index = 1
            
        elif line.startswith("### "):
            title = clean_inline(line[4:])
            script_block_index = 1
            badge = ""
            if "Workbook Exercise" in title or "📝" in title:
                badge = "[WORKBOOK EXERCISE] "
                title = re.sub(r"^(📝\s*)?(Sisterly\s*)?Workbook\s*Exercise:\s*", "", title)
            elif "Safety" in title:
                badge = "[SAFETY] "
                
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            
            if badge:
                r_badge = p.add_run(badge)
                set_run_font(r_badge, size=11.5, bold=True, color=SECONDARY_COLOR if "SAFETY" in badge else PRIMARY_COLOR)
                
            r_title = p.add_run(title)
            set_run_font(r_title, size=11.5, bold=True, color=PRIMARY_COLOR if not badge else INK)
            
            anchor = get_anchor_name(title)
            add_bookmark(p, anchor, bm_id)
            bm_id += 1
            
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.375)
            p.paragraph_format.first_line_indent = Inches(-0.188)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(clean_inline(line[2:]))
            set_run_font(r, color=INK)
            
        elif re.match(r"^(\d+)\. ", line):
            num = re.match(r"^(\d+)\. ", line).group(1)
            text_part = re.sub(r"^\d+\. ", "", line)
            
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.first_line_indent = Inches(-0.15)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.2
            
            r_num = p.add_run(f"{num}.  ")
            set_run_font(r_num, bold=True, color=PRIMARY_COLOR)
            
            r_text = p.add_run(clean_inline(text_part))
            set_run_font(r_text, color=INK)
            
        else:
            add_para(doc, clean_inline(line))
            
    if code_lines:
        if is_thought_block(code_lines):
            add_thought_quote(doc, code_lines)
        else:
            badge = None
            if current_part == "Part 4":
                if script_block_index == 1:
                    badge = "[SLIDE 1/3: SOFT & WARM]"
                elif script_block_index == 2:
                    badge = "[SLIDE 2/3: CLEAR & FIRM]"
                elif script_block_index == 3:
                    badge = "[SLIDE 3/3: PROTECTIVE]"
            elif current_part == "Part 5":
                badge = "[PUSHBACK DIALOGUE]"
            elif current_part == "Part 6":
                badge = "[QUICK START / COPY-PASTE]"
            add_script_card(doc, code_lines, badge_text=badge)
            
    if quote_lines:
        add_blockquote(doc, quote_lines)


def build():
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_value_pages(doc)
    add_markdown_remainder(doc)
    doc.save(OUT)
    print(f"Generated beautiful premium guide at: {OUT}")


if __name__ == "__main__":
    build()
